from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path('outputs/Cahier_des_charges_MVP_Security_Center_VSCode.docx')
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = '2E74B5'; DARK = '1F4D78'; NAVY = '17365D'; LIGHT = 'E8EEF5'; GRAY = 'F2F4F7'; RED='9B1C1C'; GREEN='2E7D32'

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn('w:'+m))
        if node is None: node=OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def set_repeat_table_header(row):
    trPr=row._tr.get_or_add_trPr(); el=OxmlElement('w:tblHeader'); el.set(qn('w:val'),'true'); trPr.append(el)

def set_table_widths(table, widths):
    table.autofit=False
    for row in table.rows:
        for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    tblPr=table._tbl.tblPr; tblW=tblPr.find(qn('w:tblW'))
    if tblW is None: tblW=OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'),'9360'); tblW.set(qn('w:type'),'dxa')
    tblInd=tblPr.find(qn('w:tblInd'))
    if tblInd is None: tblInd=OxmlElement('w:tblInd'); tblPr.append(tblInd)
    tblInd.set(qn('w:w'),'120'); tblInd.set(qn('w:type'),'dxa')

def add_table(doc, headers, rows, widths):
    t=doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'; set_table_widths(t,widths)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; set_cell_shading(c,GRAY); set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs: r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY)
    set_repeat_table_header(t.rows[0])
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=str(val); set_cell_margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs: p.paragraph_format.space_after=Pt(2); p.paragraph_format.line_spacing=1.05
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t

def bullet(doc, text, level=0):
    p=doc.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2'); p.add_run(text); return p

def numbered(doc,text):
    p=doc.add_paragraph(style='List Number'); p.add_run(text); return p

def callout(doc, label, text, fill=LIGHT):
    t=doc.add_table(rows=1,cols=1); set_table_widths(t,[6.5]); c=t.cell(0,0); set_cell_shading(c,fill); set_cell_margins(c,140,180,140,180)
    p=c.paragraphs[0]; r=p.add_run(label+' — '); r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY); p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)

doc=Document(); sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
sec.header_distance=sec.footer_distance=Inches(.492)
styles=doc.styles
normal=styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.10
for name,size,color,before,after in [('Title',25,NAVY,0,6),('Subtitle',13,'666666',0,16),('Heading 1',16,BLUE,16,8),('Heading 2',13,BLUE,12,6),('Heading 3',12,DARK,8,4)]:
    s=styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=(name!='Subtitle'); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
for lname in ['List Bullet','List Bullet 2','List Number']:
    styles[lname].font.name='Calibri'; styles[lname].font.size=Pt(11); styles[lname].paragraph_format.space_after=Pt(5); styles[lname].paragraph_format.line_spacing=1.10

# Header/footer
hp=sec.header.paragraphs[0]; hp.text='SECURITY CENTER • Cahier des charges MVP'; hp.style=styles['Caption']; hp.runs[0].font.color.rgb=RGBColor.from_string('777777')
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld)

# Cover
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(50); p.paragraph_format.space_after=Pt(5)
r=p.add_run('CAHIER DES CHARGES'); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor.from_string(BLUE)
p=doc.add_paragraph(style='Title'); p.add_run('Security Center for VS Code')
p=doc.add_paragraph(style='Subtitle'); p.add_run('Extension DevSecOps intégrée pour détecter, expliquer, corréler et valider les vulnérabilités')
add_table(doc,['Champ','Valeur'],[
 ('Nature du projet','Projet de fin d’études (PFA) — MVP produit'),('Produit principal','Extension Visual Studio Code installable (.vsix)'),('Utilisateurs','Développeurs, responsables sécurité et équipes DevSecOps'),('Version','1.0 — 15 juillet 2026'),('Statut','Spécification fonctionnelle et technique proposée')],[1.8,4.7])
callout(doc,'Vision','Permettre à un développeur de lancer des analyses de sécurité, comprendre les résultats, localiser les vulnérabilités dans son code et vérifier leur correction sans quitter Visual Studio Code.')
doc.add_page_break()

doc.add_heading('1. Résumé exécutif',level=1)
doc.add_paragraph("Security Center est une extension VS Code qui orchestre des scanners open source, normalise leurs résultats et les présente directement dans l’environnement de développement. Le MVP privilégie un parcours local simple : installation de l’extension, analyse du workspace, affichage des alertes dans l’onglet Problems et validation après correction. Un mode équipe, appuyé sur FastAPI et PostgreSQL, complète ce parcours avec historique, dashboard et Quality Gates.")
doc.add_paragraph("La partie pentest est strictement limitée aux applications locales ou explicitement autorisées. Elle combine un scan passif OWASP ZAP avec des scénarios contrôlés SQLi, XSS réfléchi et IDOR. Les requêtes et réponses servant de preuves sont nettoyées afin de ne pas conserver de secrets.")

doc.add_heading('2. Contexte et problématique',level=1)
doc.add_paragraph("Les outils de sécurité applicative produisent des formats, niveaux de sévérité et identifiants différents. Les développeurs doivent souvent quitter l’IDE, interpréter des rapports séparés et vérifier manuellement si une correction est effective. Cette fragmentation ralentit la remédiation et favorise les doublons, les faux positifs et les régressions.")
doc.add_heading('2.1 Problématique',level=2)
callout(doc,'Question centrale','Comment intégrer dans VS Code une chaîne DevSecOps utilisable au quotidien, capable de détecter plusieurs catégories de vulnérabilités, de présenter des preuves exploitables et de valider automatiquement les corrections ?')

doc.add_heading('3. Objectifs',level=1)
doc.add_heading('3.1 Objectif général',level=2)
doc.add_paragraph("Concevoir et réaliser une extension VS Code installable qui centralise les analyses SAST, secrets, dépendances et pentest autorisé, puis guide le développeur jusqu’à la validation de la correction.")
doc.add_heading('3.2 Objectifs spécifiques',level=2)
for x in [
 'Lancer un scan rapide ou complet depuis VS Code.', 'Afficher chaque vulnérabilité dans Problems et sur la ligne de code concernée.',
 'Orchestrer Semgrep, Gitleaks et Trivy dans le MVP.', 'Normaliser les résultats dans un modèle interne commun et générer un fingerprint stable.',
 'Dédupliquer les alertes et suivre leur état entre deux scans.', 'Expliquer la vulnérabilité et proposer des pistes de correction adaptées au contexte.',
 'Exécuter un pentest limité sur une cible autorisée et sauvegarder une preuve reproductible.', 'Rejouer le contrôle après correction et détecter les régressions.',
 'Fournir un mode équipe avec historique PostgreSQL, dashboard React et Quality Gate Jenkins.'
]: bullet(doc,x)

doc.add_heading('4. Périmètre du MVP',level=1)
add_table(doc,['Inclus dans le MVP','Reporté après le MVP'],[
 ('Extension VS Code installable, scans rapides/complets, diagnostics et panneau latéral','Support de tous les IDE et de tous les frameworks'),
 ('Semgrep, Gitleaks, Trivy ; import SARIF/JSON','OSV-Scanner avancé et scanners commerciaux multiples'),
 ('ZAP passif + SQLi/XSS/IDOR contrôlés','Pentest autonome généraliste et exploitation automatisée'),
 ('Fingerprint, déduplication simple, historique et statuts','Corrélation probabiliste complexe multi-outils'),
 ('Explication IA facultative via Ollama/API, avec validation humaine','Correction automatique sans confirmation'),
 ('FastAPI, PostgreSQL, React, Docker Compose, Jenkins','Architecture Kubernetes à haute disponibilité'),
 ('Preuves avant/après et rejeu de scénarios','Forensic complet et conservation illimitée du trafic')],[3.25,3.25])

doc.add_heading('5. Utilisateurs et rôles',level=1)
add_table(doc,['Rôle','Besoins principaux','Droits MVP'],[
 ('Développeur','Détecter, comprendre, corriger et valider','Scanner son workspace, consulter, relancer, classer un faux positif'),
 ('Security/DevSecOps','Superviser et définir les règles','Configurer scanners, règles, cibles et Quality Gates'),
 ('Lead/Manager','Suivre le risque et les tendances','Consulter dashboard et états agrégés'),
 ('Administrateur','Exploiter la plateforme','Gérer projets, utilisateurs, rétention et intégrations')],[1.25,2.75,2.5])

doc.add_heading('6. Parcours utilisateur principal',level=1)
for x in ['Installer le fichier .vsix et ouvrir un workspace.', 'Initialiser ou accepter le fichier .security-center.yml.', 'Choisir Scan rapide ou Scan complet.', 'Suivre l’avancement dans la barre d’état et le panneau Security Center.', 'Ouvrir une alerte depuis Problems à la ligne concernée.', 'Consulter la description, la sévérité, la preuve et la recommandation.', 'Modifier le code puis sélectionner Relancer ce contrôle.', 'Obtenir le statut Corrigée si le finding disparaît et, le cas échéant, si le replay échoue.', 'Synchroniser facultativement le résultat avec le serveur d’équipe.']: numbered(doc,x)

doc.add_heading('7. Exigences fonctionnelles',level=1)
reqs=[
 ('F-01','Installation et activation','L’extension est distribuée en .vsix et s’active à l’ouverture d’un workspace.','Must'),
 ('F-02','Scan rapide','Analyser les fichiers modifiés avec Semgrep et Gitleaks.','Must'),
 ('F-03','Scan complet','Analyser code, secrets, dépendances et configurations avec Semgrep, Gitleaks et Trivy.','Must'),
 ('F-04','Diagnostics VS Code','Afficher fichier, plage, règle, message et sévérité dans Problems.','Must'),
 ('F-05','Vue Security Center','Regrouper les findings par sévérité, scanner, catégorie et statut.','Must'),
 ('F-06','Navigation','Un clic ouvre le fichier à la ligne exacte.','Must'),
 ('F-07','Actions rapides','Voir détails, expliquer, faux positif, relancer et afficher preuve.','Must'),
 ('F-08','Normalisation','Transformer les sorties JSON/SARIF vers un schéma commun.','Must'),
 ('F-09','Cycle de vie','Gérer OPEN, CONFIRMED, IN_PROGRESS, FIXED, FALSE_POSITIVE, ACCEPTED_RISK.','Must'),
 ('F-10','Déduplication','Créer un fingerprint sur projet, règle, fichier, symbole et emplacement stabilisé.','Must'),
 ('F-11','Pentest autorisé','Exiger une cible allowlistée et une confirmation avant tout test actif.','Must'),
 ('F-12','Replay HTTP','Enregistrer une requête nettoyée et la rejouer après correction.','Should'),
 ('F-13','Preuve avant/après','Conserver scanners, requêtes/réponses expurgées, fichiers et résultat de validation.','Should'),
 ('F-14','IA assistée','Expliquer et proposer une correction sans l’appliquer silencieusement.','Should'),
 ('F-15','Dashboard','Filtrer findings par projet, branche, sévérité, outil, statut et date.','Should'),
 ('F-16','Quality Gate','Produire un code de sortie non nul selon la politique de sécurité.','Must'),
 ('F-17','Régression','Signaler la réapparition d’un fingerprint précédemment corrigé.','Should'),
 ('F-18','Export','Exporter un rapport JSON/SARIF et une synthèse avant/après.','Could')]
add_table(doc,['ID','Fonction','Critère synthétique','Priorité'],reqs,[.55,1.45,3.8,.7])

doc.add_heading('8. Partie pentest encadrée',level=1)
doc.add_paragraph("Le module n’est pas un outil d’exploitation libre. Il sert à confirmer une hypothèse sur une application de test possédée ou autorisée. Par défaut, seules les cibles localhost et les services Docker du projet sont acceptées.")
add_table(doc,['Test MVP','Entrée','Méthode de confirmation','Résultat'],[
 ('ZAP passif','URL ou OpenAPI','Analyse des en-têtes, cookies, CORS et réponses sans payload actif','Alerte passive et preuve'),
 ('SQLi contrôlée','Endpoint + paramètre','Baseline, valeur de contrôle, test borné, comparaison erreur/temps/contenu','Non confirmé / probable / confirmé'),
 ('XSS réfléchi','Endpoint + paramètre','Détection du reflet et analyse du contexte de sortie','Contexte et preuve expurgée'),
 ('IDOR','Deux comptes de test et deux ressources','Rejeu croisé et vérification de propriété métier','Accès refusé ou accès non autorisé')],[1.15,1.45,2.75,1.15])
doc.add_heading('8.1 Garde-fous obligatoires',level=2)
for x in ['Allowlist de domaines/cibles ; refus par défaut des IP publiques.', 'Confirmation explicite avant un test actif.', 'Limitation du débit, du nombre de requêtes et du temps d’exécution.', 'Interdiction des tests destructifs, du déni de service et de la persistance.', 'Utilisation de comptes et données de test.', 'Masquage des Authorization, cookies, tokens, mots de passe et données personnelles.', 'Journal d’audit indiquant qui a lancé quoi, quand et sur quelle cible.']: bullet(doc,x)

doc.add_heading('9. Architecture cible',level=1)
doc.add_paragraph('Architecture logique du MVP :')
for x in ['Extension VS Code : commandes, Tree View, diagnostics, actions rapides et configuration.', 'Runner local : exécution bornée des scanners et conversion des sorties.', 'API FastAPI : projets, scans, findings, preuves, politiques et validations.', 'Workers Docker : isolation de chaque scanner.', 'PostgreSQL : traçabilité des scans et du cycle de vie.', 'Dashboard React : suivi global et tendances.', 'Jenkins : déclenchement CI et application du Quality Gate.']: bullet(doc,x)
callout(doc,'Flux principal','VS Code → orchestrateur local/API → scanner Docker → normaliseur → moteur de corrélation → PostgreSQL → diagnostics VS Code et dashboard.')

doc.add_heading('10. Composants techniques',level=1)
add_table(doc,['Composant','Technologie','Responsabilités'],[
 ('Extension','TypeScript, VS Code API','UX, commandes, diagnostics, Tree View, configuration et appels locaux/API'),
 ('Backend','Python, FastAPI, Pydantic','API, orchestration, normalisation, politiques et authentification'),
 ('Base','PostgreSQL','Projets, scans, findings, occurrences, preuves, scénarios et audit'),
 ('Dashboard','React + TypeScript','Filtres, détails, tendances et historique'),
 ('Scanners','Semgrep, Gitleaks, Trivy, OWASP ZAP','SAST, secrets, SCA/config, DAST passif/contrôlé'),
 ('IA facultative','Ollama ou API LLM','Explication et proposition de correction avec contexte minimal'),
 ('Exécution','Docker Compose','Isolation et reproductibilité'),
 ('CI/CD','Jenkins','Scan automatisé, archive SARIF et Quality Gate')],[1.25,1.75,3.5])

doc.add_heading('11. Modèle de données',level=1)
add_table(doc,['Entité','Champs essentiels','Rôle'],[
 ('projects','id, name, repository_url, default_branch','Projet analysé'),('scans','id, project_id, commit_sha, branch, type, status, timestamps','Une campagne d’analyse'),('scanner_executions','scan_id, tool, version, status, duration, raw_artifact','Exécution d’un outil'),('findings','id, fingerprint, category, title, severity, status, first_seen, last_seen','Identité logique'),('finding_occurrences','finding_id, scan_id, file, lines, endpoint, evidence','Présence dans un scan'),('http_scenarios','target, method, path, parameter, auth_profile_ref, assertions','Scénario rejouable'),('validation_evidence','finding_id, before_ref, after_ref, result, timestamp','Preuve de correction'),('audit_events','actor, action, target, timestamp, metadata','Traçabilité')],[1.45,3.55,1.5])

doc.add_heading('12. API minimale',level=1)
add_table(doc,['Méthode','Route','Usage'],[
 ('POST','/api/projects','Créer ou enregistrer un projet'),('POST','/api/scans','Lancer une analyse'),('GET','/api/scans/{id}','Suivre l’état d’un scan'),('GET','/api/projects/{id}/findings','Lister et filtrer les findings'),('PATCH','/api/findings/{id}','Changer statut ou justification'),('POST','/api/findings/{id}/validate','Relancer le contrôle associé'),('POST','/api/pentest/scenarios','Créer un scénario autorisé'),('POST','/api/pentest/scenarios/{id}/replay','Rejouer un scénario'),('POST','/api/quality-gates/evaluate','Évaluer la politique'),('GET','/api/dashboard/summary','Obtenir les indicateurs')],[.8,2.6,3.1])

doc.add_heading('13. Configuration projet',level=1)
doc.add_paragraph("Un fichier .security-center.yml versionné définit les comportements reproductibles : scanners activés, exclusions, déclencheurs, politique de sévérité, cibles autorisées et paramètres de confidentialité. Les secrets et profils d’authentification ne sont jamais stockés dans ce fichier ; ils utilisent VS Code SecretStorage ou un coffre externe.")
callout(doc,'Exemple de politique','Scan rapide à la sauvegarde ; scan complet manuel/CI ; build bloqué sur CRITICAL ou sur toute vulnérabilité confirmée dynamiquement ; pentest limité à http://localhost:8080.')

doc.add_heading('14. Corrélation et cycle de vie',level=1)
for x in ['Normaliser la catégorie (CWE/OWASP), la sévérité et l’emplacement.', 'Calculer un fingerprint stable et regrouper les occurrences identiques.', 'Augmenter la confiance si un finding Semgrep est confirmé par un scénario dynamique sur le même endpoint/paramètre.', 'Conserver séparément l’identité du finding et ses occurrences par scan.', 'Passer à FIXED lorsque le finding n’est plus détecté et que le contrôle de validation réussit.', 'Passer à REGRESSION lorsqu’un finding FIXED réapparaît sur un commit ultérieur.']: numbered(doc,x)

doc.add_heading('15. Exigences non fonctionnelles',level=1)
add_table(doc,['ID','Exigence','Cible MVP'],[
 ('NF-01','Performance','Scan rapide ≤ 30 s sur le projet de démonstration ; interface non bloquante.'),('NF-02','Fiabilité','Échec isolé d’un scanner sans perte des résultats des autres.'),('NF-03','Sécurité','Moindre privilège, validation des chemins/URLs, secrets chiffrés ou SecretStorage.'),('NF-04','Confidentialité','Aucun code envoyé à un LLM distant sans consentement et configuration explicite.'),('NF-05','Compatibilité','VS Code récent sur Windows/Linux ; Docker recommandé pour scans complets.'),('NF-06','Observabilité','Logs structurés, durées, statuts et erreurs par scanner.'),('NF-07','Maintenabilité','Adaptateur indépendant par outil et tests unitaires des normaliseurs.'),('NF-08','Accessibilité','Navigation clavier, libellés lisibles et couleurs non utilisées seules.'),('NF-09','Scalabilité','Mode équipe capable de mettre les scans en file sans bloquer l’API.'),('NF-10','Auditabilité','Version de l’outil, règle, commit et configuration associés à chaque résultat.')],[.65,1.4,4.45])

doc.add_heading('16. Critères d’acceptation du MVP',level=1)
for x in ['Un développeur installe le .vsix sans modifier le code source.', 'Sur un projet volontairement vulnérable, Scan Workspace lance au moins Semgrep, Gitleaks et Trivy.', 'Les résultats apparaissent dans Problems et ouvrent le bon fichier à la bonne ligne.', 'Un second scan ne crée pas de doublon logique pour le même finding.', 'Après correction, la vulnérabilité passe à FIXED avec une preuve de validation.', 'Un scénario autorisé SQLi, XSS ou IDOR peut être enregistré et rejoué sur localhost.', 'Une cible non allowlistée est refusée.', 'Le dashboard affiche au minimum nombre de findings par sévérité, outil et statut.', 'Le pipeline Jenkins échoue lorsque le Quality Gate configuré est violé.', 'Les tokens, cookies et mots de passe sont masqués dans les logs et preuves.']: bullet(doc,'☐ '+x)

doc.add_heading('17. Stratégie de tests',level=1)
add_table(doc,['Niveau','Tests'],[
 ('Unitaires','Parseurs JSON/SARIF, fingerprints, mapping des sévérités, règles de Quality Gate, redaction des secrets.'),('Intégration','Runner ↔ scanner Docker ; API ↔ PostgreSQL ; extension ↔ API.'),('End-to-end','Installer le .vsix, scanner une application vulnérable, corriger, valider et vérifier le dashboard.'),('Sécurité','SSRF sur cible DAST, path traversal, injection de commande, fuite de secret, droits API.'),('Régression','Rejouer les scénarios validés sur les commits suivants.'),('Performance','Mesurer durée, mémoire et comportement avec plusieurs scanners.')],[1.3,5.2])

doc.add_heading('18. Planning de réalisation proposé (14 semaines)',level=1)
add_table(doc,['Semaines','Livrable'],[
 ('1–2','Architecture, backlog, maquettes, modèle normalisé et projet vulnérable de démonstration.'),('3–4','Extension minimale + Semgrep + diagnostics VS Code.'),('5','Gitleaks et gestion sécurisée des secrets.'),('6','Trivy, scan complet et agrégation des résultats.'),('7','Fingerprint, déduplication, statuts et validation statique.'),('8–9','FastAPI, PostgreSQL et synchronisation mode équipe.'),('10','Dashboard React et filtres.'),('11','ZAP passif, allowlist et garde-fous.'),('12','Un scénario contrôlé prioritaire (IDOR), puis SQLi/XSS selon avancement.'),('13','Jenkins, Quality Gate, preuve avant/après et régression.'),('14','Tests finaux, packaging .vsix, documentation, métriques et préparation soutenance.')],[1.1,5.4])

doc.add_heading('19. Backlog priorisé',level=1)
add_table(doc,['Lot','User story','Definition of Done'],[
 ('L1','En tant que développeur, je lance Semgrep depuis VS Code.','Commande disponible, progression visible, erreurs gérées.'),('L2','Je vois les résultats sur les lignes concernées.','Diagnostics, Tree View et navigation fonctionnels.'),('L3','Je lance Gitleaks et Trivy dans le scan complet.','Sorties normalisées et fusionnées.'),('L4','Je corrige puis relance un contrôle.','Statut FIXED et preuve enregistrée.'),('L5','Je synchronise avec le serveur.','API et PostgreSQL conservent scans/findings.'),('L6','Je consulte l’état global.','Dashboard filtrable et détail d’un finding.'),('L7','Je teste une cible locale autorisée.','Allowlist, ZAP passif et un scénario IDOR contrôlé.'),('L8','Je protège la CI.','Quality Gate Jenkins reproductible.'),('L9','Je demande une explication IA.','Opt-in, contexte minimal, résultat marqué comme suggestion.')],[.65,2.8,3.05])

doc.add_heading('20. Risques et mesures de maîtrise',level=1)
add_table(doc,['Risque','Impact','Mesure'],[
 ('Périmètre trop large','Retard et produit incomplet','Geler le MVP ; une attaque dynamique prioritaire avant les variantes.'),('Installation complexe','Faible adoption','Mode local simple, détection des prérequis et messages guidés.'),('Faux positifs','Perte de confiance','Preuves, niveaux de confiance, baseline et statut justifié.'),('Pentest sur cible non autorisée','Risque légal et opérationnel','Allowlist, confirmation, débit limité, audit et localhost par défaut.'),('Fuite de secrets vers IA/logs','Incident de confidentialité','Redaction, Ollama local par défaut, consentement explicite.'),('Formats outils instables','Régressions d’intégration','Adaptateurs versionnés, fixtures et tests de contrat.'),('Temps de scan long','Mauvaise UX','Scan différentiel, tâches asynchrones et annulation.')],[2.0,1.3,3.2])

doc.add_heading('21. Livrables attendus',level=1)
for x in ['Extension VS Code packagée au format .vsix.', 'Code source versionné et documenté.', 'Backend FastAPI et schéma PostgreSQL.', 'Dashboard React.', 'Docker Compose des services et scanners.', 'Pipeline Jenkins avec Quality Gate.', 'Application volontairement vulnérable et jeux de tests autorisés.', 'Documentation utilisateur, installation et architecture.', 'Rapport PFA, démonstration avant/après et résultats d’évaluation.']: bullet(doc,x)

doc.add_heading('22. Indicateurs d’évaluation',level=1)
add_table(doc,['Indicateur','Mesure proposée'],[
 ('Précision','Taux de vrais positifs sur un jeu de vulnérabilités connues.'),('Couverture','Catégories détectées par SAST, secrets, SCA/config et DAST.'),('Déduplication','Réduction du nombre d’alertes présentées après corrélation.'),('Temps de correction','Durée entre détection et validation.'),('Performance','Durée médiane par type de scan.'),('Régression','Pourcentage de scénarios rejoués avec succès en CI.'),('Utilisabilité','Temps nécessaire à un nouveau développeur pour installer et lancer le premier scan.')],[1.7,4.8])

doc.add_heading('23. Scénario de démonstration finale',level=1)
for x in ['Ouvrir dans VS Code une application de démonstration contenant une injection et un IDOR.', 'Lancer Scan complet et afficher Semgrep/Gitleaks/Trivy dans Problems.', 'Sélectionner l’IDOR et lancer le scénario local avec deux comptes de test.', 'Afficher la corrélation endpoint → contrôleur → ligne et la preuve HTTP expurgée.', 'Appliquer manuellement la correction proposée.', 'Relancer le scan et le replay ; montrer le statut FIXED.', 'Pousser le commit ; montrer Jenkins et le Quality Gate.', 'Afficher dans React l’historique avant/après et l’absence de régression.']: numbered(doc,x)

doc.add_heading('24. Hors périmètre et règles éthiques',level=1)
doc.add_paragraph("Le MVP n’effectue aucune attaque sur Internet sans autorisation, aucun déni de service, aucune persistance, aucune exfiltration et aucune exploitation destructive. Les scénarios sont conçus pour des laboratoires locaux, environnements de staging approuvés ou applications dont l’utilisateur possède une autorisation explicite. Toute fonctionnalité d’IA reste une aide à la décision : le développeur conserve la responsabilité de la correction et de sa revue.")

doc.add_heading('25. Définition de réussite du PFA',level=1)
callout(doc,'MVP réussi','Un développeur externe au projet peut installer le .vsix, analyser une application de démonstration, comprendre une vulnérabilité à la ligne concernée, corriger le code, rejouer la validation et constater la mise à jour du statut — le tout principalement depuis VS Code.',fill='E2F0D9')

# Keep table rows together and style font sizes
for t in doc.tables:
    for row in t.rows:
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.name='Calibri'; r.font.size=Pt(9.2)

doc.core_properties.title='Cahier des charges MVP — Security Center for VS Code'
doc.core_properties.subject='Extension DevSecOps intégrée à Visual Studio Code'
doc.core_properties.author='Projet PFA'
doc.save(OUT)
print(OUT.resolve())
